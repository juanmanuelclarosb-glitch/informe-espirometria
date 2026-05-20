import flet as ft
import requests
import base64
import json
import os
from docx import Document

# Clave fija de Gemini establecida
GEMINI_API_KEY = "AIzaSyCxmbJWFmelDYDdTDnyc-SPH1uDNxT_WD0"

def main(page: ft.Page):
    page.title = "Asistente de Espirometría"
    page.scroll = "adaptive"
    page.theme_mode = ft.ThemeMode.LIGHT
    
    # Variables de estado para almacenar los datos extraídos
    datos_extraidos = {}
    
    # Componentes visuales de la interfaz
    status_text = ft.Text(value="Selecciona una imagen para comenzar", color=ft.colors.BLUE_GREY)
    progreso = ft.ProgressRing(visible=False)
    
    # Campos de edición interactivos
    input_nombre = ft.TextField(label="Nombre del Paciente", expand=True)
    input_edad = ft.TextField(label="Edad", expand=True)
    input_fecha = ft.TextField(label="Fecha", expand=True)
    input_conclusion = ft.TextField(label="Conclusión Clínica", multiline=True, min_lines=3)
    
    # Contenedor del editor (oculto al principio)
    editor_container = ft.Column(visible=False, spacing=15)
    
    def analizar_imagen(e: ft.FilePickerResultEvent):
        if not e.files:
            return
        
        file_path = e.files[0].path
        status_text.value = "Leyendo imagen y conectando con Gemini..."
        progreso.visible = True
        page.update()
        
        try:
            # Convertir imagen a Base64
            with open(file_path, "rb") as image_file:
                base64_image = base64.b64encode(image_file.read()).decode('utf-8')
            
            # Preparar la petición para la API de Gemini
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={GEMINI_API_KEY}"
            
            prompt = (
                "Analiza la imagen de la espirometría. Extrae los valores numéricos y responde "
                "ÚNICAMENTE con un objeto JSON válido y limpio, sin formato markdown ni texto extra, usando exactamente estas claves: "
                "nombre, edad, fecha, farmaco, fvc_bas, fvc_bas_pct, fvc_post, fvc_post_pct, fvc_cambio, "
                "fev1_bas, fev1_bas_pct, fev1_post, fev1_post_pct, fev1_cambio, rel_bas, rel_bas_pct, "
                "rel_post, rel_post_pct, rel_cambio, conclusion"
            )
            
            payload = {
                "contents": [{
                    "parts": [
                        {"text": prompt},
                        {"inlineData": {"mimeType": "image/jpeg", "data": base64_image}}
                    ]
                }]
            }
            
            response = requests.post(url, json=payload, headers={"Content-Type": "application/json"})
            res_json = response.json()
            respuesta_texto = res_json['candidates'][0]['content']['parts'][0]['text'].strip()
            
            # Limpieza de posibles bloques markdown
            if respuesta_texto.startswith("```"):
                respuesta_texto = respuesta_texto.split("```")[1]
                if respuesta_texto.startswith("json"):
                    respuesta_texto = respuesta_texto[4:]
            
            global datos_extraidos
            datos_extraidos = json.loads(respuesta_texto.strip())
            
            # Rellenar campos del editor visual
            input_nombre.value = datos_extraidos.get("nombre", "")
            input_edad.value = datos_extraidos.get("edad", "")
            input_fecha.value = datos_extraidos.get("fecha", "")
            input_conclusion.value = datos_extraidos.get("conclusion", "")
            
            status_text.value = "¡Datos extraídos! Por favor verifícalos abajo antes de exportar."
            editor_container.visible = True
            
        except Exception as ex:
            status_text.value = f"Error en el análisis: {str(ex)}"
        
        progreso.visible = False
        page.update()

    def generar_word(e):
        status_text.value = "Generando documento Word..."
        page.update()
        
        plantilla_path = "MODELO INFORME.docx"
        if not os.path.exists(plantilla_path):
            status_text.value = "Error: No se encontró 'MODELO INFORME.docx' en la carpeta de la app."
            page.update()
            return
            
        try:
            doc = Document(plantilla_path)
            
            # Mapear las correcciones del usuario en la pantalla
            reemplazos = {
                "[NOMBRE]": input_nombre.value,
                "[EDAD]": input_edad.value,
                "[FECHA]": input_fecha.value,
                "[CONCLUSION]": input_conclusion.value,
                "[FVC_BAS]": datos_extraidos.get("fvc_bas", ""),
                "[FVC_BAS_PCT]": datos_extraidos.get("fvc_bas_pct", ""),
                "[FVC_POST]": datos_extraidos.get("fvc_post", ""),
                "[FVC_POST_PCT]": datos_extraidos.get("fvc_post_pct", ""),
                "[FVC_CAMBIO]": datos_extraidos.get("fvc_cambio", ""),
                "[FEV1_BAS]": datos_extraidos.get("fev1_bas", ""),
                "[FEV1_BAS_PCT]": datos_extraidos.get("fev1_bas_pct", ""),
                "[FEV1_POST]": datos_extraidos.get("fev1_post", ""),
                "[FEV1_POST_PCT]": datos_extraidos.get("fev1_post_pct", ""),
                "[FEV1_CAMBIO]": datos_extraidos.get("fev1_cambio", ""),
                "[REL_BAS]": datos_extraidos.get("rel_bas", ""),
                "[REL_BAS_PCT]": datos_extraidos.get("rel_bas_pct", ""),
                "[REL_POST]": datos_extraidos.get("rel_post", ""),
                "[REL_POST_PCT]": datos_extraidos.get("rel_post_pct", ""),
                "[REL_CAMBIO]": datos_extraidos.get("rel_cambio", ""),
            }
            
            # Algoritmo de reemplazo de texto en Word
            for p in doc.paragraphs:
                for codigo, valor in reemplazos.items():
                    if codigo in p.text:
                        p.text = p.text.replace(codigo, valor)
            
            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for p in celda.paragraphs:
                            for codigo, valor in reemplazos.items():
                                if codigo in p.text:
                                    p.text = p.text.replace(codigo, valor)
            
            output_name = f"Informe_{input_nombre.value.replace(' ', '_')}.docx"
            doc.save(output_name)
            status_text.value = f"¡Éxito! Archivo guardado como: {output_name}"
            
        except Exception as ex:
            status_text.value = f"Error al escribir el Word: {str(ex)}"
        page.update()

    # Configurar el selector de archivos del teléfono
    file_picker = ft.FilePicker(on_result=analizar_imagen)
    page.overlay.append(file_picker)
    
    # Construir el formulario del editor
    editor_container.controls = [
        ft.Row([input_nombre, input_edad]),
        ft.Row([input_fecha]),
        input_conclusion,
        ft.ElevatedButton(
            "Generar Documento Word", 
            icon=ft.icons.FILE_DOWNLOAD, 
            on_click=generar_word,
            style=ft.ButtonStyle(bgcolor=ft.colors.GREEN_700, color=ft.colors.WHITE)
        )
    ]

    # Diseño general de la pantalla principal
    page.add(
        ft.Card(
            content=ft.Container(
                content=ft.Column([
                    ft.Text("Asistente Clínico de Espirometrías", font_family="sans-serif", weight=ft.FontWeight.BOLD, size=18),
                    ft.Text("Sube la foto del ticket para procesar el informe automáticamente.", size=12, color=ft.colors.GREY_600),
                    ft.Divider(),
                    ft.ElevatedButton(
                        "Tomar Foto / Seleccionar Imagen",
                        icon=ft.icons.CAMERA_ALT,
                        on_click=lambda _: file_picker.pick_files(allow_multiple=False, file_type=ft.FilePickerFileType.IMAGE)
                    ),
                    ft.Row([progreso, status_text], alignment=ft.MainAxisAlignment.CENTER),
                ], alignment=ft.MainAxisAlignment.CENTER, horizontal_alignment=ft.CrossAxisAlignment.CENTER),
                padding=20
            )
        ),
        editor_container
    )

if __name__ == "__main__":
    ft.app(target=main)

